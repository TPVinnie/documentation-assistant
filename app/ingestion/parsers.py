"""File parsers for PDF, DOCX, Markdown, and TXT (FR-01, FR-02).

Each parser turns a file into a `RawDocument`: a flat list of `ContentUnit`s
(one per page for PDF, one per heading-delimited section for DOCX/Markdown,
a single body unit for plain TXT) plus whatever header metadata (title,
version, effective_date, status, supersedes) the document declares. Parsers
raise typed exceptions on failure; the ingestion pipeline decides how to
record each outcome in the processing report — parsers themselves never
swallow errors silently.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader
from pypdf.errors import PdfReadError

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}

_HEADER_KEYS = {"title", "category", "version", "effective_date", "date", "status", "supersedes"}
_HEADER_LINE_RE = re.compile(r"^\s*([A-Za-z_ ]+)\s*:\s*(.+?)\s*$")


class UnsupportedFileTypeError(ValueError):
    pass


class EmptyDocumentError(ValueError):
    pass


class EncryptedDocumentError(ValueError):
    pass


class CorruptedDocumentError(ValueError):
    pass


@dataclass
class ContentUnit:
    """One citable, indexable slice of a document (a page or a section)."""

    index: int
    label: str
    text: str


@dataclass
class RawDocument:
    file_path: Path
    file_name: str
    file_type: str  # pdf | docx | md | txt
    header: dict[str, str] = field(default_factory=dict)
    units: list[ContentUnit] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n\n".join(u.text for u in self.units)


def _extract_front_matter(lines: list[str]) -> tuple[dict[str, str], int]:
    """Parses leading `key: value` lines (optionally fenced by `---`) into a dict.

    Returns the header dict and the index of the first line of body content.
    """
    header: dict[str, str] = {}
    start = 0
    fenced = bool(lines) and lines[0].strip() == "---"
    if fenced:
        start = 1
    i = start
    while i < len(lines):
        line = lines[i]
        if fenced and line.strip() == "---":
            i += 1
            break
        if not line.strip():
            if not fenced:
                break
            i += 1
            continue
        match = _HEADER_LINE_RE.match(line)
        if not match:
            break
        key = match.group(1).strip().lower().replace(" ", "_")
        if key not in _HEADER_KEYS:
            break
        header[key] = match.group(2).strip()
        i += 1
    if "date" in header and "effective_date" not in header:
        header["effective_date"] = header["date"]
    return header, i


def parse_txt(path: Path) -> RawDocument:
    raw = path.read_text(encoding="utf-8", errors="replace")
    if not raw.strip():
        raise EmptyDocumentError(f"{path.name} has no extractable text")
    lines = raw.splitlines()
    header, body_start = _extract_front_matter(lines)
    body = "\n".join(lines[body_start:]).strip()
    if not body:
        raise EmptyDocumentError(f"{path.name} has no body text after header")
    units = [ContentUnit(index=0, label="document body", text=body)]
    return RawDocument(file_path=path, file_name=path.name, file_type="txt", header=header, units=units)


_MD_HEADER_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def parse_markdown(path: Path) -> RawDocument:
    raw = path.read_text(encoding="utf-8", errors="replace")
    if not raw.strip():
        raise EmptyDocumentError(f"{path.name} has no extractable text")
    lines = raw.splitlines()
    header, body_start = _extract_front_matter(lines)

    units: list[ContentUnit] = []
    current_label = "document body"
    current_lines: list[str] = []
    unit_index = 0

    def flush() -> None:
        nonlocal unit_index
        text = "\n".join(current_lines).strip()
        if text:
            units.append(ContentUnit(index=unit_index, label=current_label, text=text))
            unit_index += 1

    for line in lines[body_start:]:
        heading = _MD_HEADER_RE.match(line)
        if heading:
            flush()
            current_lines = []
            current_label = heading.group(2).strip()
        else:
            current_lines.append(line)
    flush()

    if not units:
        raise EmptyDocumentError(f"{path.name} has no body text after header")
    return RawDocument(file_path=path, file_name=path.name, file_type="md", header=header, units=units)


def parse_docx(path: Path) -> RawDocument:
    try:
        doc = DocxDocument(str(path))
    except Exception as exc:  # python-docx raises bare Exception/PackageNotFoundError on corrupt files
        raise CorruptedDocumentError(f"{path.name} could not be opened: {exc}") from exc

    paragraphs = [p for p in doc.paragraphs]
    header: dict[str, str] = {}
    header_line_count = 0
    header_lines = [p.text for p in paragraphs[:8]]
    parsed_header, consumed = _extract_front_matter(header_lines)
    if parsed_header:
        header = parsed_header
        header_line_count = consumed

    units: list[ContentUnit] = []
    current_label = "document body"
    current_lines: list[str] = []
    unit_index = 0

    def flush() -> None:
        nonlocal unit_index
        text = "\n".join(current_lines).strip()
        if text:
            units.append(ContentUnit(index=unit_index, label=current_label, text=text))
            unit_index += 1

    for p in paragraphs[header_line_count:]:
        style = (p.style.name if p.style is not None else "") or ""
        if style.lower().startswith("heading") and p.text.strip():
            flush()
            current_lines = []
            current_label = p.text.strip()
        elif p.text.strip():
            current_lines.append(p.text)
    flush()

    if not units:
        raise EmptyDocumentError(f"{path.name} has no extractable text")
    return RawDocument(file_path=path, file_name=path.name, file_type="docx", header=header, units=units)


def parse_pdf(path: Path) -> RawDocument:
    try:
        reader = PdfReader(str(path))
    except PdfReadError as exc:
        raise CorruptedDocumentError(f"{path.name} could not be read: {exc}") from exc

    if reader.is_encrypted:
        try:
            result = reader.decrypt("")
        except Exception:
            result = 0
        if not result:
            raise EncryptedDocumentError(f"{path.name} is password-protected")

    units: list[ContentUnit] = []
    header: dict[str, str] = {}
    for i, page in enumerate(reader.pages):
        try:
            text = (page.extract_text() or "").strip()
        except Exception as exc:
            raise CorruptedDocumentError(f"{path.name} page {i + 1} failed to extract: {exc}") from exc

        if i == 0 and text:
            # A blank line in the source doesn't survive PDF text extraction (no text
            # object marks it), so front matter here can't rely on a blank-line/`---`
            # separator the way parse_markdown/parse_txt do — body_start just lands on
            # the first line that isn't a recognized `key: value` header line.
            lines = text.splitlines()
            header, body_start = _extract_front_matter(lines)
            text = "\n".join(lines[body_start:]).strip()

        if text:
            units.append(ContentUnit(index=i, label=f"page {i + 1}", text=text))

    if not units:
        raise EmptyDocumentError(f"{path.name} has no extractable text (possibly scanned/image-only)")
    return RawDocument(file_path=path, file_name=path.name, file_type="pdf", header=header, units=units)


_PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".md": parse_markdown,
    ".txt": parse_txt,
}


def parse_document(path: Path) -> RawDocument:
    suffix = path.suffix.lower()
    parser = _PARSERS.get(suffix)
    if parser is None:
        raise UnsupportedFileTypeError(f"{path.name} has unsupported extension '{suffix}'")
    return parser(path)
